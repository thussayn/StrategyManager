# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import streamlit as st

SESSION_LOGO_KEY = "_logo_path"
AR_FONT_FAMILY = "Cairo"

def get_logo_path(assets_dir: str | None = None):
    p = st.session_state.get(SESSION_LOGO_KEY)
    if p and os.path.exists(p):
        return p
    if assets_dir:
        lp = os.path.join(assets_dir, "Full Logo.png")
        if os.path.exists(lp):
            return lp
    return None

def _run_set_ar_font(run, bold=False):
    try:
        run.font.name = AR_FONT_FAMILY
        run.font.bold = bool(bold)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), AR_FONT_FAMILY)
        rFonts.set(qn('w:hAnsi'), AR_FONT_FAMILY)
        rFonts.set(qn('w:cs'), AR_FONT_FAMILY)
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:bidi'), 'ar-SA')
    except Exception:
        pass

def _rtl_paragraph(p):
    try:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        bidi.set(qn('w:val'), '1')
        for r in p.runs:
            _run_set_ar_font(r)
    except Exception:
        pass

def _add_heading_rtl(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph("", style=style)
    _rtl_paragraph(p)
    r = p.add_run(text)
    _run_set_ar_font(r, bold=True)
    return p

def _add_paragraph_rtl(doc, text, style=None):
    p = doc.add_paragraph("", style=style) if style else doc.add_paragraph("")
    _rtl_paragraph(p)
    r = p.add_run(text)
    _run_set_ar_font(r, bold=False)
    return p

def _add_bullet_rtl(doc, text):
    p = doc.add_paragraph("", style="List Bullet")
    _rtl_paragraph(p)
    r = p.add_run(text)
    _run_set_ar_font(r, bold=False)
    return p

def _ensure_doc_defaults_rtl(doc):
    try:
        styles_elm = doc.styles.element
        docDefaults = styles_elm.find(qn('w:docDefaults'))
        if docDefaults is None:
            docDefaults = OxmlElement('w:docDefaults')
            styles_elm.append(docDefaults)
        pPrDefault = docDefaults.find(qn('w:pPrDefault'))
        if pPrDefault is None:
            pPrDefault = OxmlElement('w:pPrDefault')
            docDefaults.append(pPrDefault)
        pPr = pPrDefault.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            pPrDefault.append(pPr)
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        bidi.set(qn('w:val'), '1')
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'both')
    except Exception:
        pass

def export_to_word(strategy: dict, assets_dir: str | None = None):
    doc = Document()
    _ensure_doc_defaults_rtl(doc)
    for st_name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
        try:
            style = doc.styles[st_name]
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            pass

    lp = get_logo_path(assets_dir)
    if lp and os.path.exists(lp):
        try:
            doc.add_picture(lp, width=__import__('docx').shared.Inches(1.8))
        except Exception:
            pass

    _add_heading_rtl(doc, str(strategy.get("name","")), level=1)
    _add_heading_rtl(doc, "الرؤية:", level=1)
    _add_paragraph_rtl(doc, str(strategy.get("vision","")))
    _add_heading_rtl(doc, "الرسالة:", level=1)
    _add_paragraph_rtl(doc, str(strategy.get("message","")))
    _add_heading_rtl(doc, "الأهداف:", level=1)
    for g in strategy.get("goals", []):
        _add_bullet_rtl(doc, str(g))
    _add_heading_rtl(doc, "القيم:", level=1)
    for v in strategy.get("values", []):
        _add_bullet_rtl(doc, str(v))

    for p in doc.paragraphs:
        _rtl_paragraph(p)

    fname = f"strategy_{strategy.get('id','')}.docx"
    out_path = os.path.join(os.getcwd(), fname)
    doc.save(out_path)
    return out_path
